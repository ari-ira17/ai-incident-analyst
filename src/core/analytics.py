import polars as pl
import os
import json
import ollama
import pandas as pd
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class IncidentAnalytics:
    def __init__(self, df: pl.DataFrame):
        self.df = df

    def get_top_regions(self, limit: int = 10) -> pl.DataFrame:
        """Возвращает датафрейм с топ-N регионами по сумме баллов severity"""
        return (
            self.df.group_by("Муниципалитет")
            .agg(pl.col("severity").sum().alias("Суммы баллов"))
            .sort("Суммы баллов", descending=True)
            .head(limit)
        )

    def get_region_distribution(self, region: str) -> pl.DataFrame:
        """Возвращает распределение проблем (в штуках и %) для конкретного региона"""
        region_df = self.df.filter(pl.col("Муниципалитет") == region)
        total_incidents = region_df.height
        
        if total_incidents == 0:
            return pl.DataFrame()
            
        return (
            region_df.group_by("Тема")
            .agg(pl.len().alias("count"))
            .with_columns(
                (pl.col("count") / total_incidents * 100).alias("percentage")
            )
            .sort("count", descending=True)
        )

    def get_top_problems_with_quotes(self, region: str, top_n: int = 3, quotes_per_problem: int = 2) -> list:
        """Собирает топ-проблемы и по N цитат для формирования промпта в Ollama"""
        dist_df = self.get_region_distribution(region).head(top_n)
        result = []
        
        if dist_df.is_empty():
            return result
            
        for row in dist_df.iter_rows():
            theme = row[0]
            percentage = row[2]
            
            quotes = (
                self.df.filter(
                    (pl.col("Муниципалитет") == region) & 
                    (pl.col("Тема") == theme) &
                    (pl.col("Текст инцидента").is_not_null())
                )
                .select("Текст инцидента")
                .head(quotes_per_problem)
                .to_series()
                .to_list()
            )
            
            result.append({
                "theme": theme,
                "percentage": percentage,
                "quotes": quotes
            })
            
        return result

    def _prepare_excel_data(self, top_df: pl.DataFrame) -> pl.DataFrame:
        """Вспомогательный метод для формирования структуры листов Excel"""
        rows = []
        for idx, row in enumerate(top_df.iter_rows(named=True), 1):
            mun_name = row["Муниципалитет"]
            score = row["Суммы баллов"]
            
            dist = self.get_region_distribution(mun_name).head(3)
            if not dist.is_empty():
                top_themes = dist["Тема"].to_list()
                themes_str = ", ".join(top_themes)
            else:
                themes_str = "Нет данных"
                
            rows.append({
                "Место": idx,
                "Муниципалитет": mun_name,
                "Суммы баллов": score,
                "Проблемы": themes_str
            })
        return pl.DataFrame(rows)

    def build_reports(self, docx_output_path: str, xlsx_output_path: str):
        """Основной метод генерации документов Excel и текстового отчета Word через Ollama"""
        import docx  
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        top_3_raw = self.get_top_regions(3)
        top_10_raw = self.get_top_regions(10)
        
        df_excel_top3 = self._prepare_excel_data(top_3_raw)
        df_excel_top10 = self._prepare_excel_data(top_10_raw)
        
        # 1. Генерация Excel (оставляем без изменений)
        os.makedirs(os.path.dirname(xlsx_output_path), exist_ok=True)
        with pd.ExcelWriter(xlsx_output_path, engine="xlsxwriter") as writer:
            df_excel_top3.to_pandas().to_excel(writer, sheet_name="Топ-3", index=False)
            df_excel_top10.to_pandas().to_excel(writer, sheet_name="Топ-10", index=False)
            
        # Подготовка контекста для LLM
        ai_prompt_context = "АНАЛИТИЧЕСКИЕ ДАННЫЕ ДЛЯ АНАЛИЗА\n===================================\n\n"
        
        for idx, row in enumerate(top_10_raw.iter_rows(named=True), 1):
            mun_name = row["Муниципалитет"]
            score = row["Суммы баллов"]
            
            ai_prompt_context += f"{idx}. Муниципалитет: {mun_name} (Сумма баллов критичности: {score})\n"
            problems_data = self.get_top_problems_with_quotes(mun_name, top_n=3, quotes_per_problem=2)
            
            for p in problems_data:
                ai_prompt_context += f"   - Проблема: '{p['theme']}' ({p['percentage']:.1f}% от всех инцидентов региона)\n"
                for q_idx, quote in enumerate(p["quotes"], 1):
                    clean_quote = str(quote).strip().replace('\n', ' ')
                    ai_prompt_context += f"     Цитата {q_idx}: \"{clean_quote}\"\n"
            ai_prompt_context += "\n"

        system_instruction = (
            "Ты — ведущий аналитик государственного ситуационного центра. Твоя задача — составить официальный отчет.\n\n"
            "СТРОГИЕ ПРАВИЛА ОФОРМЛЕНИЯ:\n"
            "1. Отчет должен быть строго разбит на блоки по муниципалитетам.\n"
            "2. Каждый блок ОБЯЗАТЕЛЬНО начинается с названия муниципалитета (например: '1. Москаленский район').\n"
            "3. Сразу под названием муниципалитета должны идти ровно 3 его главные проблемы в виде списка через дефис (-).\n"
            "4. Категорически запрещены любые вводные фразы, приветствия, пояснения (вроде 'Вот упрощенная версия данных:') или подзаголовки.\n"
            "5. Никакого Markdown: полностью запрещено использовать решетки (#), звездочки (*) или жирные выделения (**).\n"
            "6. Пиши строго по схеме:\n"
            "Название муниципалитета\n"
            "- Проблема 1\n"
            "- Проблема 2\n"
            "- Проблема 3"
        )
        
        # Инструкция по формату, которая подмешивается в prompt для фиксации паттерна ответа
        format_instruction = (
            "ПИШИ СТРОГО ПО ЭТОМУ ШАБЛОНУ ДЛЯ КАЖДОГО РЕГИОНА ИЗ СПИСКА:\n"
            "1. [Сюда подставь название муниципалитета из данных]\n"
            "- [Первая проблема с кратким анализом]\n"
            "- [Вторая проблема с кратким анализом]\n"
            "- [Третья проблема с кратким анализом]\n"
            "2. [Сюда подставь название следующего муниципалитета]\n"
            "...и так далее. Никакого другого текста, кроме названий и 3-х дефисов под ними, быть не должно."
        )

        print("Запрос к локальной модели Ollama...")
        try:
            response = ollama.generate(
                model="qwen2.5:3b", 
                system=system_instruction,
                prompt=f"{format_instruction}\n\nИСХОДНЫЕ ДАННЫЕ ДЛЯ ОБРАБОТКИ:\n{ai_prompt_context}"
            )
            final_report_text = response["response"]
        except Exception as e:
            final_report_text = f"Ошибка выполнения запроса к Ollama: {str(e)}\n\nСгенерированные данные:\n{ai_prompt_context}"
            
        # 2. ГЕНЕРАЦИЯ ОФИЦИАЛЬНОГО WORD-ДОКУМЕНТА (.docx)
        os.makedirs(os.path.dirname(docx_output_path), exist_ok=True)
        
        doc = docx.Document()
        
        # Настройка глобального стиля Normal (основной текст)
        style_normal = doc.styles['Normal']
        font_normal = style_normal.font
        font_normal.name = 'Times New Roman'
        font_normal.size = Pt(14)
        font_normal.color.rgb = RGBColor(0, 0, 0)

        # --- ГЛАВНЫЙ ЗАГОЛОВОК (Level 1) ---
        title_text = "АНАЛИТИЧЕСКИЙ ОТЧЕТ СИТУАЦИОННОГО ЦЕНТРА"
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_after = Pt(12)
        
        run_title = p_title.add_run(title_text)
        run_title.font.name = 'Times New Roman'
        run_title.font.size = Pt(16)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(0, 0, 0)

        # --- ПОДЗАГОЛОВОК ДОКУМЕНТА ---
        subtitle_text = "Мониторинг системных инцидентов и структуры проблем в разрезе муниципалитетов"
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(24)
        
        run_sub = p_sub.add_run(subtitle_text)
        run_sub.font.name = 'Times New Roman'
        run_sub.font.size = Pt(12)
        run_sub.font.italic = True
        run_sub.font.color.rgb = RGBColor(0, 0, 0)
        
        # --- НАПОЛНЕНИЕ ТЕКСТОМ С ЖЕСТКОЙ ФИЛЬТРАЦИЕЙ ---
        for paragraph_text in final_report_text.split('\n'):
            clean_text = paragraph_text.strip()
            
            # 1. Жесткий блек-лист фраз, которые мы вырезаем, если ИИ их выдал
            blacklisted_phrases = [
                "вот упрощенная", 
                "сгруппированная версия", 
                "группировка по муниципалитетам",
                "ниже представлен",
                "аналитический отчет по следующей"
            ]
            if any(phrase in clean_text.lower() for phrase in blacklisted_phrases):
                continue
                
            # 2. Программная очистка от маркдаун-мусора (решетки и звездочки)
            clean_text = clean_text.replace('#', '').replace('*', '').strip()
            
            if clean_text:
                p = doc.add_paragraph()
                
                # Проверяем, является ли строка заголовком муниципалитета (например: "1. Москаленский район")
                # Если строка начинается с цифры или содержит слово "район"/"муниципалитет" без дефиса в начале
                is_heading = (
                    any(keyword in clean_text.lower() for keyword in ["муниципалитет", "район", "вывод", "заключение"]) 
                    and not clean_text.startswith("-")
                )
                
                if is_heading:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(14)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.keep_with_next = True
                    
                    run = p.add_run(clean_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.bold = True  # Делаем заголовок района красивым и жирным, без решеток
                    run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    # Обычный абзац текста или пункт списка
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                    
                    # Если это пункт списка (начинается с дефиса), то красную строку делать не нужно
                    if clean_text.startswith("-"):
                        p.paragraph_format.first_line_indent = Pt(0)
                    else:
                        p.paragraph_format.first_line_indent = Pt(36) # Красная строка для обычного текста
                    
                    run = p.add_run(clean_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    
        # Получаем текущую дату в формате: ДД.ММ.ГГГГ
        current_date_str = datetime.now().strftime("%d.%m.%Y")
        
        # Добавляем пустую строку-отступ перед датой
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(24)
        
        # Создаем абзац для даты с выравниванием по правому краю
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run_date = p_date.add_run(f"Дата формирования отчета: {current_date_str}")
        run_date.font.name = 'Times New Roman'
        run_date.font.size = Pt(12)
        run_date.font.italic = True
        run_date.font.color.rgb = RGBColor(0, 0, 0)
        
        # Сохраняем итоговый Word-файл
        doc.save(docx_output_path)
            
        # Возвращаем результат работы метода
        return top_10_raw["Муниципалитет"].to_list()
